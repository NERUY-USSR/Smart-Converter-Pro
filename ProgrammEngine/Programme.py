import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from PIL import Image
import os
import threading
import subprocess
from pathlib import Path

# ==========================================
# ЛОГИКА ОБРАБОТКИ ФАЙЛОВ (Да, тут оформление)
# ==========================================
class FileProcessor:
    @staticmethod
    def convert_media(fin, fout):
        """Надежная конвертация аудио/видео через ffmpeg"""
        import imageio_ffmpeg as ffmpeg
        exe = ffmpeg.get_ffmpeg_exe()

        #-y (перезапись), -i (вход), выход
        cmd = [exe, "-y", "-i", fin, fout]

        #Скрываем окно консоли на Windows
        si = None
        if os.name == 'nt':
            si = subprocess.STARTUPINFO()
            si.dwFlags |= subprocess.STARTF_USESHOWWINDOW  # Исправленная константа

        result = subprocess.run(cmd, capture_output=True, startupinfo=si)
        if result.returncode != 0:
            error_msg = result.stderr.decode('utf-8', errors='ignore')
            raise RuntimeError(f"Ошибка кодеков: {error_msg}")

    @staticmethod
    def convert_text(fin, fout, target_fmt):
        """чтение PDF и DOCX"""
        from docx import Document
        content = ""
        ext_in = os.path.splitext(fin)[1].lower()
        #--- ЧТЕНИЕ ---
        try:
            if ext_in == ".pdf":
                import PyPDF2
                with open(fin, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    # Собираем текст по страницам
                    pages_text = [p.extract_text() for p in reader.pages if p.extract_text()]
                    content = "\n".join(pages_text)  # Вот здесь были пустые скобки
            elif ext_in == ".docx":
                from docx import Document
                doc = Document(fin)
                full_text = []
                #Собираем текст из всех абзацев
                for p in doc.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text)
                #Собираем текст из всех таблиц
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text)
                content = "\n".join(full_text)
            else:
                with open(fin, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()
        except Exception as e:
            raise RuntimeError(f"Ошибка при чтении файла: {str(e)}")

        if not content.strip():
            raise ValueError("Файл пуст или защищен от чтения")

        #--- ЗАПИСЬ ---
        target_fmt = target_fmt.upper()
        if target_fmt == "PDF":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            font_path = "C:/Windows/Fonts/arial.ttf"
            if os.path.exists(font_path):
                pdf.add_font('ArialRus', '', font_path)
                pdf.set_font('ArialRus', size=12)
            else:
                pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, txt=content)
            pdf.output(fout)
        elif target_fmt == "DOCX":
            new_doc = Document()
            new_doc.add_paragraph(content)
            new_doc.save(fout)
        else:  #TXT
            with open(fout, 'w', encoding='utf-8') as f:
                f.write(content)


# ==========================================
# ИНТЕРФЕЙС ПРИЛОЖЕНИЯ
# ==========================================
class SmartConverterPro:
    def __init__(self, root):
        self.root = root
        self.root.title("Smart Converter Pro (Fixed Edition)")
        self.root.geometry("550x550")

        self.format_map = {
            "Изображения": ['.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff'],
            "Тексты": ['.pdf', '.txt', '.docx'],
            "Видео": ['.mp4', '.avi', '.mkv', '.mov'],
            "Аудио": ['.mp3', '.wav', '.ogg', '.m4a', '.flac']
        }
        self.target_options = {
            "Изображения": ["PNG", "JPG", "WEBP"],
            "Тексты": ["TXT", "PDF", "DOCX"],
            "Видео": ["MP4", "AVI", "MKV"],
            "Аудио": ["MP3", "WAV", "M4A"]
        }
        self._category = None
        self.setup_ui()

    def setup_ui(self):
        #Блок "Откуда"
        tk.Label(self.root, text="Папка с файлами (Источник):", font=("Arial", 9, "bold")).pack(pady=(15, 0))
        f1 = tk.Frame(self.root);
        f1.pack(fill="x", padx=40)
        self.in_var = tk.StringVar(value=os.getcwd())
        tk.Entry(f1, textvariable=self.in_var).pack(side="left", expand=True, fill="x")
        tk.Button(f1, text="...", command=self.select_in_dir).pack(side="right")

        #Блок Названия и списка
        tk.Label(self.root, text="Выберите или введите имя файла:", font=("Arial", 9, "bold")).pack(pady=(15, 5))
        self.name_var = tk.StringVar()
        self.entry = tk.Entry(self.root, textvariable=self.name_var, width=55)
        self.entry.pack()
        self.name_var.trace_add("write", self.update_list)

        self.listbox = tk.Listbox(self.root, height=6, width=55)
        self.listbox.pack(pady=5)
        self.listbox.bind("<<ListboxSelect>>", self.on_item_select)

        #Блок "Формата"
        tk.Label(self.root, text="Конвертировать в:").pack(pady=5)
        self.fmt_combo = ttk.Combobox(self.root, state="readonly", width=20)
        self.fmt_combo.pack()

        #Блок "Куда"
        tk.Label(self.root, text="Папка сохранения:", font=("Arial", 9, "bold")).pack(pady=(15, 0))
        f2 = tk.Frame(self.root);
        f2.pack(fill="x", padx=40)
        self.out_var = tk.StringVar(value=os.path.join(Path.home(), "Desktop", "Converted_Files"))
        tk.Entry(f2, textvariable=self.out_var).pack(side="left", expand=True, fill="x")
        tk.Button(f2, text="...", command=self.select_out_dir).pack(side="right")

        #Кнопка пуск конвертации
        self.btn_start = tk.Button(self.root, text="ЗАПУСТИТЬ КОНВЕРТАЦИЮ", bg="#2ecc71", fg="white",
                                   font=("Arial", 10, "bold"), height=2, width=30, command=self.start_process)
        self.btn_start.pack(pady=25)

    def select_in_dir(self):
        d = filedialog.askdirectory()
        if d: self.in_var.set(d); self.update_list()

    def select_out_dir(self):
        d = filedialog.askdirectory()
        if d: self.out_var.set(d)

    def update_list(self, *args):
        search = self.name_var.get().lower()
        self.listbox.delete(0, tk.END)
        path = self.in_var.get()
        if os.path.exists(path):
            try:
                files = [f for f in os.listdir(path) if os.path.isfile(os.path.join(path, f))]
                for f in files:
                    if search in f.lower(): self.listbox.insert(tk.END, f)
            except:
                pass

    def on_item_select(self, event):
        if not self.listbox.curselection(): return
        name = self.listbox.get(self.listbox.curselection())
        self.name_var.set(name)

        #Получаем расширение текущего файла (например, ".m4a")
        current_ext = os.path.splitext(name)[1].lower()

        for cat, exts in self.format_map.items():
            if current_ext in exts:
                #ФИЛЬТР: оставляем только те форматы, которые НЕ совпадают с текущим
                #Мы убираем точку из current_ext (например, ".m4a" -> "m4a") и сравниваем
                filtered_opts = [fmt for fmt in self.target_options[cat]
                                 if fmt.lower() != current_ext.replace('.', '').lower()]

                self.fmt_combo['values'] = filtered_opts
                if filtered_opts:
                    self.fmt_combo.current(0)
                self._category = cat
                return
    def start_process(self):
        fin = os.path.join(self.in_var.get(), self.name_var.get())
        out_dir = self.out_var.get()
        fmt = self.fmt_combo.get()

        if not os.path.exists(fin) or not fmt:
            return messagebox.showwarning("Внимание", "Выберите файл и формат!")

        if not os.path.exists(out_dir): os.makedirs(out_dir)

        base_name = os.path.splitext(self.name_var.get())[0]
        fout = os.path.join(out_dir, f"{base_name}.{fmt.lower()}")

        #Чтобы интерфейс не зависал, запускаем в потоке (я в моменте...)
        self.btn_start.config(state="disabled", text="ОБРАБОТКА...")
        threading.Thread(target=self.worker, args=(fin, fout, fmt), daemon=True).start()

    def worker(self, fin, fout, fmt):
        try:
            if self._category in ["Видео", "Аудио"]:
                FileProcessor.convert_media(fin, fout)
            elif self._category == "Тексты":
                FileProcessor.convert_text(fin, fout, fmt)
            elif self._category == "Изображения":
                img = Image.open(fin)
                if fmt == "JPG": img = img.convert("RGB")
                img.save(fout)

            messagebox.showinfo("Успех", f"Файл успешно сохранен:\n{fout}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка:\n{str(e)}")
        finally:
            self.root.after(0, lambda: self.btn_start.config(state="normal", text="ЗАПУСТИТЬ КОНВЕРТАЦИЮ"))


if __name__ == "__main__":
    root = tk.Tk()
    app = SmartConverterPro(root)
    root.mainloop()